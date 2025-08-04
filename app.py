#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Web-приложение EPUB Cutter
Современное веб-интерфейс для создания сайтов из EPUB книг
"""

import os
import json
import sqlite3
import zipfile
from pathlib import Path
from datetime import datetime
from flask import Flask, render_template, request, jsonify, send_from_directory, redirect, url_for, flash, session, send_file
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.utils import secure_filename
import sys
import tempfile
import io

# Функции для работы с метаданными (перенесены из metadata_utils)
def extract_epub_metadata(opf_root):
    """Извлекает метаданные книги из OPF файла"""
    import xml.etree.ElementTree as ET
    
    # Пространства имен для OPF
    namespaces = {
        'opf': 'http://www.idpf.org/2007/opf',
        'dc': 'http://purl.org/dc/elements/1.1/'
    }
    
    # Извлечение названия книги
    title_elem = opf_root.find('.//dc:title', namespaces)
    title = title_elem.text if title_elem is not None and title_elem.text else "Без названия"
    
    # Извлечение автора
    author_elem = opf_root.find('.//dc:creator', namespaces)
    author = author_elem.text if author_elem is not None and author_elem.text else "Неизвестный автор"
    
    return title.strip(), author.strip()

def format_book_info(title, author):
    """Форматирует информацию о книге для отображения"""
    return f"{title} - {author}"

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-change-in-production'
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['BOOKS_FOLDER'] = 'books'
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB max file size

# Настройка Flask-Login
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'
login_manager.login_message = 'Пожалуйста, войдите в систему для доступа к этой странице.'
login_manager.login_message_category = 'info'

# Учетные данные пользователя
USERNAME = 'reading'
PASSWORD = 'readingbooks'

class User(UserMixin):
    def __init__(self, username):
        self.id = username

@login_manager.user_loader
def load_user(user_id):
    if user_id == USERNAME:
        return User(user_id)
    return None

# Убеждаемся, что папки существуют
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['BOOKS_FOLDER'], exist_ok=True)

class EPUBProcessor:
    """Класс для обработки EPUB файлов (адаптирован из основного приложения)"""
    
    def __init__(self):
        self.book_title = ""
        self.book_author = ""
        self.chapters = []
        self.images = {}
    
    def load_epub(self, epub_path):
        """Загружает EPUB файл и извлекает главы"""
        try:
            self.chapters.clear()
            self.images.clear()
            
            with zipfile.ZipFile(epub_path, 'r') as epub_zip:
                # Чтение структуры EPUB
                from xml.etree import ElementTree as ET
                
                container_xml = epub_zip.read('META-INF/container.xml')
                container_root = ET.fromstring(container_xml)
                
                opf_path = container_root.find('.//{urn:oasis:names:tc:opendocument:xmlns:container}rootfile').get('full-path')
                opf_content = epub_zip.read(opf_path)
                opf_root = ET.fromstring(opf_content)
                
                # Извлечение метаданных книги
                self.book_title, self.book_author = extract_epub_metadata(opf_root)
                
                # Загрузка изображений
                opf_dir = os.path.dirname(opf_path)
                for item in opf_root.findall('.//{http://www.idpf.org/2007/opf}item'):
                    media_type = item.get('media-type', '')
                    if media_type.startswith('image/'):
                        href = item.get('href')
                        if href:
                            img_path = os.path.join(opf_dir, href).replace('\\', '/')
                            try:
                                img_data = epub_zip.read(img_path)
                                self.images[img_path] = img_data
                            except:
                                pass
                
                # Загрузка глав
                spine_items = opf_root.findall('.//{http://www.idpf.org/2007/opf}itemref')
                manifest_items = {item.get('id'): item for item in opf_root.findall('.//{http://www.idpf.org/2007/opf}item')}
                
                # Создаем маппинг файлов для будущих ссылок
                self.file_mapping = {}
                
                for i, spine_item in enumerate(spine_items):
                    idref = spine_item.get('idref')
                    if idref in manifest_items:
                        manifest_item = manifest_items[idref]
                        href = manifest_item.get('href')
                        media_type = manifest_item.get('media-type')
                        
                        if media_type == 'application/xhtml+xml':
                            file_path = os.path.join(opf_dir, href).replace('\\', '/')
                            
                            try:
                                chapter_content = epub_zip.read(file_path)
                                content_str = chapter_content.decode('utf-8')
                                
                                # Сначала пытаемся извлечь заголовок из raw текста
                                import re
                                title_match = re.search(r'<title[^>]*>([^<]+)</title>', content_str, re.IGNORECASE)
                                title = title_match.group(1).strip() if title_match else f"Глава {i+1}"
                                
                                # Пытаемся исправить некорректные HTML комментарии для корректного XML
                                try:
                                    # Исправляем проблемные комментарии
                                    fixed_content = re.sub(r'<!--\[endif\]---->', '<!--[endif]-->', content_str)
                                    fixed_content = re.sub(r'<!--([^>]*?)---->', r'<!--\1-->', fixed_content)
                                    chapter_root = ET.fromstring(fixed_content.encode('utf-8'))
                                    
                                    # Если XML парсится успешно, используем исправленный контент
                                    chapter_content = fixed_content.encode('utf-8')
                                except:
                                    # Если исправление не помогло, используем оригинальный контент
                                    pass
                                
                                # Дополнительная попытка найти заголовок в заголовках h1-h6, если title пустой
                                if title == f"Глава {i+1}":
                                    try:
                                        if 'chapter_root' in locals():
                                            for tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                                                heading = chapter_root.find(f'.//{{{http://www.w3.org/1999/xhtml}}}{tag}')
                                                if heading is not None and heading.text:
                                                    title = heading.text.strip()
                                                    break
                                        else:
                                            # Ищем заголовки в raw тексте
                                            for tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                                                h_match = re.search(f'<{tag}[^>]*>([^<]+)</{tag}>', content_str, re.IGNORECASE)
                                                if h_match:
                                                    title = h_match.group(1).strip()
                                                    break
                                    except:
                                        pass
                                
                                # Очищаем заголовок от недопустимых символов
                                title = re.sub(r'[<>:"/\\|?*]', '', title)
                                
                                self.chapters.append({
                                    'title': title,
                                    'file_path': file_path,
                                    'content': chapter_content
                                })
                                
                                # Сохраняем маппинг для обработки ссылок
                                self.file_mapping[file_path] = len(self.chapters) - 1
                                
                            except Exception as e:
                                print(f"Ошибка при чтении главы {file_path}: {e}")
                                # Даже если возникла критическая ошибка, сохраняем главу
                                try:
                                    raw_content = epub_zip.read(file_path)
                                    title = f"Глава {i+1}"
                                    
                                    self.chapters.append({
                                        'title': title,
                                        'file_path': file_path,
                                        'content': raw_content
                                    })
                                    
                                    # Сохраняем маппинг для обработки ссылок
                                    self.file_mapping[file_path] = len(self.chapters) - 1
                                    print(f"Добавлена глава с дефолтным названием: {title}")
                                except:
                                    print(f"Не удалось загрузить главу {file_path}")
                                    continue
            
            return True
            
        except Exception as e:
            print(f"Ошибка при загрузке EPUB: {e}")
            return False
    
    def create_website(self, output_path):
        """Создает веб-сайт из загруженной книги"""
        try:
            import re
            import html
            
            # Создание главной папки сайта
            site_name = re.sub(r'[<>:"/\\|?*]', '', self.book_title)
            site_path = Path(output_path) / site_name
            site_path.mkdir(parents=True, exist_ok=True)
            
            # Создание папки для изображений
            images_path = site_path / "images"
            images_path.mkdir(exist_ok=True)
            
            # CSS стили (из основного приложения)
            css_content = self._get_website_css()
            with open(site_path / "styles.css", 'w', encoding='utf-8') as f:
                f.write(css_content)
            
            # Создание index.html
            self._create_index_html(site_path)
            
            # Создание страниц глав с навигацией
            selected_chapters = [(i, chapter) for i, chapter in enumerate(self.chapters)]
            
            for idx, (i, chapter) in enumerate(selected_chapters):
                chapter_num = i + 1
                safe_title = re.sub(r'[<>:"/\\|?*]', '', chapter['title'])
                filename = f"chapter_{chapter_num:02d}_{safe_title}.html"
                filepath = site_path / filename
                
                # Обработка содержимого главы
                content = chapter['content'].decode('utf-8')
                content = content.replace('xmlns="http://www.w3.org/1999/xhtml"', '')
                
                # Обработка изображений
                content = self._process_images_for_website(content, images_path)
                
                # Обработка внутренних ссылок
                content = self._process_internal_links(content, selected_chapters, idx)
                
                # Создание навигации
                prev_link = ""
                next_link = ""
                
                if idx > 0:
                    prev_chapter = selected_chapters[idx-1][1]
                    prev_safe_title = re.sub(r'[<>:"/\\|?*]', '', prev_chapter['title'])
                    prev_num = selected_chapters[idx-1][0] + 1
                    prev_filename = f"chapter_{prev_num:02d}_{prev_safe_title}.html"
                    prev_link = f'<a href="{prev_filename}" class="nav-button">← Предыдущая</a>'
                
                if idx < len(selected_chapters) - 1:
                    next_chapter = selected_chapters[idx+1][1]
                    next_safe_title = re.sub(r'[<>:"/\\|?*]', '', next_chapter['title'])
                    next_num = selected_chapters[idx+1][0] + 1
                    next_filename = f"chapter_{next_num:02d}_{next_safe_title}.html"
                    next_link = f'<a href="{next_filename}" class="nav-button">Следующая →</a>'
                
                # Создание полного HTML документа главы
                chapter_html = f"""<!DOCTYPE html>
<html lang="ru" class="theme-vintage">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{html.escape(chapter['title'])} - {html.escape(self.book_title)}</title>
    <link rel="stylesheet" href="styles.css">
</head>
<body class="theme-vintage">
    <!-- Переключатель тем -->
    <div class="theme-switcher">
        <div class="theme-switcher-label">🎨 Тема</div>
        <div class="theme-buttons">
            <div class="theme-btn theme-btn-vintage active" onclick="setTheme('vintage')" title="Винтажная тема"></div>
            <div class="theme-btn theme-btn-dark" onclick="setTheme('dark')" title="Черная тема"></div>
        </div>
    </div>
    
    <div class="book-header">
        <h1 class="book-title">{html.escape(self.book_title)}</h1>
        <p class="book-author">{html.escape(self.book_author)}</p>
    </div>
    
    <div class="navigation">
        <a href="index.html" class="nav-button">📚 Содержание</a>
        {prev_link}
        {next_link}
    </div>
    
    <div class="download-section">
        <h3>📥 Скачать главу:</h3>
        <div class="download-buttons">
            <button onclick="downloadChapter('epub')" class="download-btn epub-btn">📖 EPUB</button>
            <button onclick="downloadChapter('docx')" class="download-btn docx-btn">📄 DOCX</button>
        </div>
    </div>
    
    <div class="chapter-content" data-chapter-title="{html.escape(chapter['title'])}">
        {content}
    </div>
    
    <div class="navigation">
        <a href="index.html" class="nav-button">📚 Содержание</a>
        {prev_link}
        {next_link}
    </div>
    
    <script src="/static/js/notes.js"></script>
    <script>
        // Функция переключения тем
        function setTheme(theme) {{
            // Удаляем все классы тем с body и html
            document.body.className = document.body.className.replace(/theme-\\w+/g, '');
            document.documentElement.className = document.documentElement.className.replace(/theme-\\w+/g, '');
            
            // Добавляем новый класс темы к body и html
            document.body.classList.add('theme-' + theme);
            document.documentElement.classList.add('theme-' + theme);
            
            // Сохраняем выбор в localStorage
            localStorage.setItem('reading-theme', theme);
            
            // Обновляем активную кнопку
            document.querySelectorAll('.theme-btn').forEach(btn => btn.classList.remove('active'));
            document.querySelector('.theme-btn-' + theme).classList.add('active');
        }}
        
        // Загружаем сохраненную тему при загрузке страницы
        function loadSavedTheme() {{
            const savedTheme = localStorage.getItem('reading-theme') || 'vintage';
            setTheme(savedTheme);
        }}
        
        // Функция скачивания главы
        function downloadChapter(format) {{
            const pathParts = window.location.pathname.split('/');
            if (pathParts[1] === 'book' && pathParts[2]) {{
                const bookPath = pathParts[2];
                
                // Сначала получаем book_id
                fetch(`/api/book-info?path=${{encodeURIComponent(bookPath)}}`)
                    .then(response => response.json())
                    .then(data => {{
                        if (data.book_id) {{
                            const chapterIndex = {i};  // Индекс текущей главы
                            const downloadUrl = `/download-chapter/${{data.book_id}}/${{chapterIndex}}/${{format}}`;
                            
                            // Создаем временную ссылку для скачивания
                            const link = document.createElement('a');
                            link.href = downloadUrl;
                            link.style.display = 'none';
                            document.body.appendChild(link);
                            link.click();
                            document.body.removeChild(link);
                        }} else {{
                            alert('Не удалось получить информацию о книге');
                        }}
                    }})
                    .catch(error => {{
                        console.error('Ошибка:', error);
                        alert('Ошибка при скачивании файла');
                    }});
            }}
        }}
        
        // Устанавливаем информацию о книге для системы заметок
        document.addEventListener('DOMContentLoaded', function() {{
            // Загружаем сохраненную тему
            loadSavedTheme();
            
            // Получаем book_id из пути URL
            const pathParts = window.location.pathname.split('/');
            if (pathParts[1] === 'book' && pathParts[2]) {{
                const bookPath = pathParts[2];
                fetch(`/api/book-info?path=${{encodeURIComponent(bookPath)}}`)
                    .then(response => response.json())
                    .then(data => {{
                        if (data.book_id && window.notesSystem) {{
                            window.notesSystem.bookId = data.book_id;
                            window.notesSystem.addNotesButtonToNavigation();
                        }}
                    }})
                    .catch(error => console.log('Could not load book info:', error));
            }}
        }});
    </script>
</body>
</html>"""
                
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(chapter_html)
            
            return str(site_path)
            
        except Exception as e:
            print(f"Ошибка при создании сайта: {e}")
            return None
    
    def _create_index_html(self, site_path):
        """Создает главную страницу с оглавлением"""
        import html
        import re
        
        selected_chapters = [(i, chapter) for i, chapter in enumerate(self.chapters)]
        
        # Создаем словарь соответствий для ссылок в оглавлении
        self._create_chapter_mapping(selected_chapters)
        
        html_content = f"""<!DOCTYPE html>
<html lang="ru" class="theme-vintage">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{self.book_title}</title>
    <link rel="stylesheet" href="styles.css">
</head>
<body class="theme-vintage">
    <!-- Переключатель тем -->
    <div class="theme-switcher">
        <div class="theme-switcher-label">🎨 Тема</div>
        <div class="theme-buttons">
            <div class="theme-btn theme-btn-vintage active" onclick="setTheme('vintage')" title="Винтажная тема"></div>
            <div class="theme-btn theme-btn-dark" onclick="setTheme('dark')" title="Черная тема"></div>
        </div>
    </div>
    
    <div class="book-header">
        <h1 class="book-title">{html.escape(self.book_title)}</h1>
        <p class="book-author">{html.escape(self.book_author)}</p>
    </div>
    
    <div class="toc">
        <h2>📚 Содержание</h2>
        <ul>
"""
        
        for i, (chapter_idx, chapter) in enumerate(selected_chapters):
            chapter_num = chapter_idx + 1
            safe_title = re.sub(r'[<>:"/\\|?*]', '', chapter['title'])
            filename = f"chapter_{chapter_num:02d}_{safe_title}.html"
            
            html_content += f"""            <li>
                <a href="{filename}">
                    <span>{html.escape(chapter['title'])}</span>
                    <span class="chapter-number">{chapter_num}</span>
                </a>
            </li>
"""
        
        html_content += f"""        </ul>
    </div>
    
    <div class="navigation">
        <p>Эта книга содержит {len(selected_chapters)} глав. Выберите главу из оглавления выше для чтения.</p>
    </div>
    
    <script>
        // Функция переключения тем
        function setTheme(theme) {{
            // Удаляем все классы тем с body и html
            document.body.className = document.body.className.replace(/theme-\\w+/g, '');
            document.documentElement.className = document.documentElement.className.replace(/theme-\\w+/g, '');
            
            // Добавляем новый класс темы к body и html
            document.body.classList.add('theme-' + theme);
            document.documentElement.classList.add('theme-' + theme);
            
            // Сохраняем выбор в localStorage
            localStorage.setItem('reading-theme', theme);
            
            // Обновляем активную кнопку
            document.querySelectorAll('.theme-btn').forEach(btn => btn.classList.remove('active'));
            document.querySelector('.theme-btn-' + theme).classList.add('active');
        }}
        
        // Загружаем сохраненную тему при загрузке страницы
        function loadSavedTheme() {{
            const savedTheme = localStorage.getItem('reading-theme') || 'vintage';
            setTheme(savedTheme);
        }}
        
        // Функция скачивания главы по индексу
        function downloadChapterByIndex(chapterIndex, format) {{
            const pathParts = window.location.pathname.split('/');
            if (pathParts[1] === 'book' && pathParts[2]) {{
                const bookPath = pathParts[2];
                
                fetch(`/api/book-info?path=${{encodeURIComponent(bookPath)}}`)
                    .then(response => response.json())
                    .then(data => {{
                        if (data.book_id) {{
                            const downloadUrl = `/download-chapter/${{data.book_id}}/${{chapterIndex}}/${{format}}`;
                            
                            const link = document.createElement('a');
                            link.href = downloadUrl;
                            link.style.display = 'none';
                            document.body.appendChild(link);
                            link.click();
                            document.body.removeChild(link);
                        }} else {{
                            alert('Не удалось получить информацию о книге');
                        }}
                    }})
                    .catch(error => {{
                        console.error('Ошибка:', error);
                        alert('Ошибка при скачивании файла');
                    }});
            }}
        }}
        
        // Добавляем кнопки скачивания к каждой главе
        document.addEventListener('DOMContentLoaded', function() {{
            // Загружаем сохраненную тему
            loadSavedTheme();
            
            const chapters = document.querySelectorAll('.toc li');
            chapters.forEach((li, index) => {{
                const downloadDiv = document.createElement('div');
                downloadDiv.className = 'chapter-download-buttons';
                downloadDiv.innerHTML = `
                    <button onclick="downloadChapterByIndex(${{index}}, 'epub')" class="mini-download-btn epub-btn" title="Скачать в EPUB">📖</button>
                    <button onclick="downloadChapterByIndex(${{index}}, 'docx')" class="mini-download-btn docx-btn" title="Скачать в DOCX">📄</button>
                `;
                li.appendChild(downloadDiv);
            }});
        }});
    </script>
</body>
</html>"""
        
        with open(site_path / "index.html", 'w', encoding='utf-8') as f:
            f.write(html_content)
    
    def _process_images_for_website(self, content, images_path):
        """Обрабатывает изображения для веб-сайта"""
        import re
        
        img_pattern = r'<img[^>]*src=["\']([^"\']+)["\'][^>]*>'
        
        def replace_img(match):
            img_tag = match.group(0)
            src = match.group(1)
            src_clean = src.replace('../', '').replace('./', '')
            
            # Ищем изображение по имени файла
            for img_path, img_data in self.images.items():
                img_filename = os.path.basename(img_path)
                src_filename = os.path.basename(src_clean)
                
                if img_filename == src_filename:
                    # Сохраняем изображение в папку images
                    img_filepath = images_path / img_filename
                    with open(img_filepath, 'wb') as f:
                        f.write(img_data)
                    
                    # Заменяем путь в HTML
                    relative_path = f"images/{img_filename}"
                    new_img_tag = img_tag.replace(f'src="{src}"', f'src="{relative_path}"')
                    new_img_tag = new_img_tag.replace(f"src='{src}'", f'src="{relative_path}"')
                    return new_img_tag
                    
            return img_tag
        
        return re.sub(img_pattern, replace_img, content)
    
    def _process_internal_links(self, content, selected_chapters, current_idx):
        """Обрабатывает внутренние ссылки между главами"""
        import re
        
        # Создаем словарь соответствий: исходный файл -> новый файл
        link_mapping = {}
        
        for idx, (i, chapter) in enumerate(selected_chapters):
            original_path = chapter['file_path']
            # Получаем имя исходного файла без пути
            original_filename = original_path.split('/')[-1]
            
            # Создаем имя нового файла
            chapter_num = i + 1
            safe_title = re.sub(r'[<>:"/\\|?*]', '', chapter['title'])
            new_filename = f"chapter_{chapter_num:02d}_{safe_title}.html"
            
            # Добавляем различные варианты ссылок
            link_mapping[original_filename] = new_filename
            link_mapping[original_path] = new_filename
            
            # Также обрабатываем относительные пути
            if '../' in original_path:
                relative_path = original_path.replace('../', '')
                link_mapping[relative_path] = new_filename
            
            # Обрабатываем пути без расширения (часто используется в ссылках)
            name_without_ext = original_filename.replace('.xhtml', '').replace('.html', '')
            link_mapping[name_without_ext] = new_filename
            
            # Также добавляем вариант с точкой в начале (относительный путь)
            if not original_filename.startswith('./'):
                link_mapping['./' + original_filename] = new_filename
        
        # Паттерн для поиска ссылок
        link_patterns = [
            r'<a[^>]*href=["\'](.*?)["\'][^>]*>',  # обычные ссылки
            r'href=["\'](.*?)["\']',  # просто href атрибуты
        ]
        
        def replace_link(match):
            full_match = match.group(0)
            href_value = match.group(1)
            
            # Пропускаем внешние ссылки и якоря
            if href_value.startswith(('http:', 'https:', 'mailto:', '#')):
                return full_match
            
            # Разделяем ссылку на файл и якорь
            if '#' in href_value:
                file_part, anchor_part = href_value.split('#', 1)
                anchor = '#' + anchor_part
            else:
                file_part = href_value
                anchor = ''
            
            # Убираем относительные пути
            clean_file = file_part.replace('../', '').replace('./', '')
            
            # Ищем соответствие в нашем словаре
            new_filename = None
            
            # Прямое совпадение
            if clean_file in link_mapping:
                new_filename = link_mapping[clean_file]
            else:
                # Поиск по частичному совпадению имени файла
                for original, new in link_mapping.items():
                    if clean_file in original or original.endswith(clean_file):
                        new_filename = new
                        break
            
            if new_filename:
                new_href = new_filename + anchor
                # Отладочный вывод (можно убрать в продакшене)
                if file_part != new_filename:
                    print(f"Заменяем ссылку: {href_value} -> {new_href}")
                return full_match.replace(href_value, new_href)
            
            return full_match
        
        # Применяем замены для всех паттернов
        for pattern in link_patterns:
            content = re.sub(pattern, replace_link, content)
        
        return content
    
    def export_chapter_to_docx(self, chapter_index, book_title, book_author):
        """Экспортирует главу в формат DOCX"""
        try:
            from docx import Document
            from docx.shared import Inches
            import re
            from html import unescape
            import xml.etree.ElementTree as ET
            
            if chapter_index >= len(self.chapters):
                return None
                
            chapter = self.chapters[chapter_index]
            
            # Создаем новый документ
            doc = Document()
            
            # Добавляем заголовок книги
            title = doc.add_heading(book_title, 0)
            author = doc.add_paragraph(f'Автор: {book_author}')
            author.alignment = 1  # Центрирование
            
            # Добавляем заголовок главы
            chapter_title = doc.add_heading(chapter['title'], level=1)
            
            # Обрабатываем содержимое главы
            content = chapter['content']
            if isinstance(content, bytes):
                content = content.decode('utf-8')
            
            # Парсим HTML для извлечения текста
            try:
                # Исправляем комментарии если нужно
                content = re.sub(r'<!--\[endif\]---->', '<!--[endif]-->', content)
                content = re.sub(r'<!--([^>]*?)---->', r'<!--\1-->', content)
                
                root = ET.fromstring(content)
                body = root.find('.//{http://www.w3.org/1999/xhtml}body')
                
                if body is not None:
                    self._process_html_to_docx(body, doc)
                else:
                    # Если не удалось парсить, извлекаем текст регулярными выражениями
                    self._extract_text_to_docx(content, doc)
                    
            except:
                # Fallback - извлекаем текст простыми методами
                self._extract_text_to_docx(content, doc)
            
            # Сохраняем в память
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            
            return docx_buffer
            
        except Exception as e:
            print(f"Ошибка при экспорте в DOCX: {e}")
            return None
    
    def _process_html_to_docx(self, element, doc):
        """Обрабатывает HTML элементы и добавляет в DOCX документ"""
        for child in element:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            
            if tag in ['p']:
                if child.text or len(child) > 0:
                    paragraph = doc.add_paragraph()
                    self._add_text_with_formatting(child, paragraph)
            elif tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(tag[1]) + 1  # h1 -> level 2, h2 -> level 3, etc.
                if child.text:
                    doc.add_heading(child.text.strip(), level=level)
            elif tag == 'div':
                self._process_html_to_docx(child, doc)
    
    def _add_text_with_formatting(self, element, paragraph):
        """Добавляет текст с форматированием в параграф"""
        if element.text:
            paragraph.add_run(element.text)
        
        for child in element:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            
            if tag == 'em' or tag == 'i':
                run = paragraph.add_run(child.text or '')
                run.italic = True
            elif tag == 'strong' or tag == 'b':
                run = paragraph.add_run(child.text or '')
                run.bold = True
            else:
                paragraph.add_run(child.text or '')
            
            if child.tail:
                paragraph.add_run(child.tail)
    
    def _extract_text_to_docx(self, html_content, doc):
        """Извлекает текст из HTML простыми методами"""
        import re
        from html import unescape
        
        # Удаляем скрипты и стили
        html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Извлекаем заголовки
        headings = re.findall(r'<h[1-6][^>]*>(.*?)</h[1-6]>', html_content, re.DOTALL | re.IGNORECASE)
        for heading in headings:
            clean_heading = re.sub(r'<[^>]+>', '', heading)
            clean_heading = unescape(clean_heading).strip()
            if clean_heading:
                doc.add_heading(clean_heading, level=2)
        
        # Извлекаем параграфы
        paragraphs = re.findall(r'<p[^>]*>(.*?)</p>', html_content, re.DOTALL | re.IGNORECASE)
        for para in paragraphs:
            clean_para = re.sub(r'<[^>]+>', '', para)
            clean_para = unescape(clean_para).strip()
            if clean_para and len(clean_para) > 10:  # Игнорируем очень короткие параграфы
                doc.add_paragraph(clean_para)
    
    def export_chapter_to_epub(self, chapter_index, book_title, book_author):
        """Экспортирует главу в формат EPUB"""
        try:
            if chapter_index >= len(self.chapters):
                return None
                
            chapter = self.chapters[chapter_index]
            
            # Создаем временный EPUB файл
            epub_buffer = io.BytesIO()
            
            with zipfile.ZipFile(epub_buffer, 'w', zipfile.ZIP_DEFLATED) as epub_zip:
                # mimetype (должен быть первым и несжатым)
                epub_zip.writestr('mimetype', 'application/epub+zip', zipfile.ZIP_STORED)
                
                # META-INF/container.xml
                container_xml = '''<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
    <rootfiles>
        <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
    </rootfiles>
</container>'''
                epub_zip.writestr('META-INF/container.xml', container_xml)
                
                # OEBPS/content.opf
                import uuid
                book_id = str(uuid.uuid4())
                
                opf_content = f'''<?xml version="1.0" encoding="UTF-8"?>
<package xmlns="http://www.idpf.org/2007/opf" unique-identifier="BookId" version="2.0">
    <metadata xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:opf="http://www.idpf.org/2007/opf">
        <dc:title>{book_title} - {chapter['title']}</dc:title>
        <dc:creator>{book_author}</dc:creator>
        <dc:identifier id="BookId">{book_id}</dc:identifier>
        <dc:language>ru</dc:language>
        <meta name="cover" content="cover"/>
    </metadata>
    <manifest>
        <item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>
        <item id="chapter" href="chapter.xhtml" media-type="application/xhtml+xml"/>
        <item id="style" href="style.css" media-type="text/css"/>
    </manifest>
    <spine toc="ncx">
        <itemref idref="chapter"/>
    </spine>
</package>'''
                epub_zip.writestr('OEBPS/content.opf', opf_content)
                
                # OEBPS/toc.ncx
                toc_content = f'''<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE ncx PUBLIC "-//NISO//DTD ncx 2005-1//EN" "http://www.daisy.org/z3986/2005/ncx-2005-1.dtd">
<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">
    <head>
        <meta name="dtb:uid" content="{book_id}"/>
        <meta name="dtb:depth" content="1"/>
        <meta name="dtb:totalPageCount" content="0"/>
        <meta name="dtb:maxPageNumber" content="0"/>
    </head>
    <docTitle>
        <text>{book_title} - {chapter['title']}</text>
    </docTitle>
    <navMap>
        <navPoint id="navpoint-1" playOrder="1">
            <navLabel>
                <text>{chapter['title']}</text>
            </navLabel>
            <content src="chapter.xhtml"/>
        </navPoint>
    </navMap>
</ncx>'''
                epub_zip.writestr('OEBPS/toc.ncx', toc_content)
                
                # OEBPS/style.css
                css_content = '''
body {
    font-family: Georgia, serif;
    line-height: 1.6;
    margin: 2em;
}
h1, h2, h3, h4, h5, h6 {
    color: #333;
    margin-top: 1.5em;
    margin-bottom: 0.5em;
}
p {
    text-align: justify;
    margin-bottom: 1em;
}
img {
    max-width: 100%;
    height: auto;
}
'''
                epub_zip.writestr('OEBPS/style.css', css_content)
                
                # OEBPS/chapter.xhtml
                content = chapter['content']
                if isinstance(content, bytes):
                    content = content.decode('utf-8')
                
                # Создаем валидный XHTML
                chapter_xhtml = f'''<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.1//EN" "http://www.w3.org/TR/xhtml11/DTD/xhtml11.dtd">
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
    <title>{chapter['title']}</title>
    <link rel="stylesheet" type="text/css" href="style.css"/>
</head>
<body>
    <h1>{chapter['title']}</h1>
    {self._clean_html_for_epub(content)}
</body>
</html>'''
                epub_zip.writestr('OEBPS/chapter.xhtml', chapter_xhtml)
            
            epub_buffer.seek(0)
            return epub_buffer
            
        except Exception as e:
            print(f"Ошибка при экспорте в EPUB: {e}")
            return None
    
    def _clean_html_for_epub(self, html_content):
        """Очищает HTML для валидного EPUB"""
        import re
        
        # Исправляем комментарии
        html_content = re.sub(r'<!--\[endif\]---->', '<!--[endif]-->', html_content)
        html_content = re.sub(r'<!--([^>]*?)---->', r'<!--\1-->', html_content)
        
        # Извлекаем содержимое body
        body_match = re.search(r'<body[^>]*>(.*?)</body>', html_content, re.DOTALL | re.IGNORECASE)
        if body_match:
            content = body_match.group(1)
        else:
            content = html_content
        
        # Убираем некорректные атрибуты и теги
        content = re.sub(r'data-[^=]*="[^"]*"', '', content)  # Удаляем data- атрибуты
        content = re.sub(r'class="[^"]*"', '', content)  # Упрощаем, убираем классы
        content = re.sub(r'style="[^"]*"', '', content)  # Убираем inline стили
        
        return content.strip()
    
    def _create_chapter_mapping(self, selected_chapters):
        """Создает маппинг исходных файлов на новые имена файлов"""
        import re
        self.chapter_mapping = {}
        
        for idx, (i, chapter) in enumerate(selected_chapters):
            original_path = chapter['file_path']
            original_filename = original_path.split('/')[-1]
            
            chapter_num = i + 1
            safe_title = re.sub(r'[<>:"/\\|?*]', '', chapter['title'])
            new_filename = f"chapter_{chapter_num:02d}_{safe_title}.html"
            
            # Различные варианты ссылок
            self.chapter_mapping[original_filename] = new_filename
            self.chapter_mapping[original_path] = new_filename
            
            if '../' in original_path:
                relative_path = original_path.replace('../', '')
                self.chapter_mapping[relative_path] = new_filename
            
            name_without_ext = original_filename.replace('.xhtml', '').replace('.html', '')
            self.chapter_mapping[name_without_ext] = new_filename
    
    def _get_website_css(self):
        """Возвращает CSS стили для сайта с поддержкой тем"""
        return """
        /* CSS переменные для тем чтения - Винтажная тема по умолчанию */
        :root {
            --bg-color: #EDE5D3;
            --text-color: #3D2914;
            --heading-color: #2A1810;
            --border-color: #B8A082;
            --card-bg: #F4EDE2;
            --button-bg: #8B4513;
            --button-hover-bg: #A0522D;
            --shadow: rgba(61, 41, 20, 0.2);
            --link-color: #8B4513;
            --accent-color: #CD853F;
            --highlight-bg: #F5DEB3;
            --muted-text: #6B4E3D;
        }

        /* Винтажная тема */
        html.theme-vintage,
        body.theme-vintage {
            --bg-color: #EDE5D3;
            --text-color: #3D2914;
            --heading-color: #2A1810;
            --border-color: #B8A082;
            --card-bg: #F4EDE2;
            --button-bg: #8B4513;
            --button-hover-bg: #A0522D;
            --shadow: rgba(61, 41, 20, 0.2);
            --link-color: #8B4513;
            --accent-color: #CD853F;
            --highlight-bg: #F5DEB3;
            --muted-text: #6B4E3D;
        }

        /* Темная тема */
        html.theme-dark,
        body.theme-dark {
            --bg-color: #000000;
            --text-color: #C0C0C0;
            --heading-color: #E0E0E0;
            --border-color: #404040;
            --card-bg: #1A1A1A;
            --button-bg: #333333;
            --button-hover-bg: #4A4A4A;
            --shadow: rgba(255, 255, 255, 0.1);
            --link-color: #A0A0A0;
            --accent-color: #666666;
            --highlight-bg: #2A2A2A;
            --muted-text: #808080;
        }

        html {
            background-color: var(--bg-color);
            transition: background-color 0.3s ease;
        }

        body {
            font-family: Georgia, "Times New Roman", serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background-color: var(--bg-color);
            color: var(--text-color);
            transition: background-color 0.3s ease, color 0.3s ease;
        }
        
        .book-header {
            text-align: center;
            border-bottom: 2px solid var(--button-bg);
            padding-bottom: 20px;
            margin-bottom: 30px;
        }
        
        .book-title {
            font-size: 2.5em;
            color: var(--heading-color);
            margin-bottom: 10px;
        }
        
        .book-author {
            font-size: 1.3em;
            color: var(--muted-text);
            font-style: italic;
        }
        
        .toc {
            background: var(--card-bg);
            padding: 20px;
            border-radius: 8px;
            border-left: 4px solid var(--accent-color);
            box-shadow: 0 2px 4px var(--shadow);
            border: 1px solid var(--border-color);
        }
        
        .toc h2 {
            color: var(--heading-color);
            margin-top: 0;
        }
        
        .toc ul {
            list-style: none;
            padding: 0;
        }
        
        .toc li {
            padding: 8px 0;
            border-bottom: 1px solid var(--border-color);
            position: relative;
        }
        
        .toc li:last-child {
            border-bottom: none;
        }
        
        .toc a {
            text-decoration: none;
            color: var(--text-color);
            font-weight: 500;
            display: block;
            padding-right: 80px;
            transition: all 0.2s ease;
        }
        
        .toc a:hover {
            color: var(--link-color);
            background-color: var(--card-bg);
            padding: 5px 10px;
            border-radius: 4px;
            margin: -5px -10px;
            margin-right: -90px;
            box-shadow: 0 1px 3px var(--shadow);
        }
        
        .chapter-download-buttons {
            position: absolute;
            right: 0;
            top: 50%;
            transform: translateY(-50%);
            display: flex;
            gap: 5px;
        }
        
        .mini-download-btn {
            padding: 4px 8px;
            border: none;
            border-radius: 3px;
            font-size: 0.8rem;
            cursor: pointer;
            transition: all 0.2s ease;
            width: 28px;
            height: 28px;
            display: flex;
            align-items: center;
            justify-content: center;
        }
        
        .mini-download-btn:hover {
            transform: scale(1.1);
        }
        
        .mini-download-btn.epub-btn {
            background: var(--button-bg);
            color: var(--card-bg);
            border: 1px solid var(--accent-color);
        }
        
        .mini-download-btn.docx-btn {
            background: var(--accent-color);
            color: var(--card-bg);
            border: 1px solid var(--button-bg);
        }
        
        .chapter-number {
            background: var(--accent-color);
            color: var(--card-bg);
            padding: 2px 8px;
            border-radius: 12px;
            font-size: 0.8em;
            min-width: 20px;
            text-align: center;
            font-weight: bold;
        }
        
        .chapter-content {
            margin-top: 30px;
            padding: 20px;
            background: var(--card-bg);
            border-radius: 8px;
            box-shadow: 0 2px 4px var(--shadow);
        }
        
        .navigation {
            margin: 30px 0;
            text-align: center;
        }
        
        .nav-button {
            background: var(--button-bg);
            color: white;
            padding: 10px 20px;
            text-decoration: none;
            border-radius: 5px;
            margin: 0 10px;
            display: inline-block;
            transition: background-color 0.2s ease;
        }
        
        .nav-button:hover {
            background: var(--button-hover-bg);
        }
        
        /* Переключатель тем */
        .theme-switcher {
            position: fixed;
            top: 20px;
            right: 20px;
            z-index: 1000;
            background: var(--card-bg);
            border-radius: 8px;
            padding: 10px;
            box-shadow: 0 2px 10px var(--shadow);
            border: 1px solid var(--border-color);
        }
        
        .theme-switcher-label {
            font-size: 0.9rem;
            color: var(--text-color);
            margin-bottom: 8px;
            display: block;
            text-align: center;
        }
        
        .theme-buttons {
            display: flex;
            gap: 5px;
        }
        
        .theme-btn {
            width: 30px;
            height: 30px;
            border: 2px solid var(--border-color);
            border-radius: 50%;
            cursor: pointer;
            transition: all 0.2s ease;
            position: relative;
        }
        
        .theme-btn:hover {
            transform: scale(1.1);
        }
        
        .theme-btn.active {
            border-color: var(--button-bg);
            box-shadow: 0 0 0 2px var(--button-bg);
        }
        
        .theme-btn-vintage { background: #EDE5D3; border-color: #B8A082; }
        .theme-btn-dark { background: #000000; border-color: #404040; }
        
        .theme-btn::after {
            content: '';
            position: absolute;
            top: 50%;
            left: 50%;
            transform: translate(-50%, -50%);
            width: 8px;
            height: 8px;
            border-radius: 50%;
        }
        
        .theme-btn-vintage::after { background: #3D2914; }
        .theme-btn-dark::after { background: #C0C0C0; }
        
        .download-section {
            background: var(--card-bg);
            padding: 20px;
            border-radius: 8px;
            margin: 20px 0;
            text-align: center;
            border-left: 4px solid var(--button-bg);
            box-shadow: 0 2px 4px var(--shadow);
        }
        
        .download-section h3 {
            color: var(--heading-color);
            margin-bottom: 15px;
            font-size: 1.2rem;
        }
        
        .download-buttons {
            display: flex;
            gap: 15px;
            justify-content: center;
            flex-wrap: wrap;
        }
        
        .download-btn {
            padding: 10px 20px;
            border: none;
            border-radius: 5px;
            font-size: 1rem;
            font-weight: 500;
            cursor: pointer;
            transition: all 0.3s ease;
            text-decoration: none;
            display: inline-flex;
            align-items: center;
            gap: 8px;
        }
        
        .epub-btn {
            background: var(--button-bg);
            color: var(--card-bg);
            border: 1px solid var(--accent-color);
        }
        
        .epub-btn:hover {
            background: var(--button-hover-bg);
            transform: translateY(-2px);
            box-shadow: 0 4px 8px var(--shadow);
        }
        
        .docx-btn {
            background: var(--accent-color);
            color: var(--card-bg);
            border: 1px solid var(--button-bg);
        }
        
        .docx-btn:hover {
            background: var(--button-bg);
            transform: translateY(-2px);
            box-shadow: 0 4px 8px var(--shadow);
        }
        
        img {
            max-width: 100%;
            height: auto;
            display: block;
            margin: 20px auto;
            border-radius: 4px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        }
        
        @media (max-width: 600px) {
            body {
                padding: 10px;
            }
            
            .book-title {
                font-size: 2em;
            }
            
            .nav-button {
                display: block;
                margin: 10px 0;
            }
            
            .download-buttons {
                flex-direction: column;
                align-items: center;
            }
            
            .download-btn {
                width: 200px;
            }
            
            .theme-switcher {
                top: 10px;
                right: 10px;
                padding: 8px;
            }
            
            .theme-btn {
                width: 25px;
                height: 25px;
            }
            
            .theme-switcher-label {
                font-size: 0.8rem;
            }
        }
        """

# Инициализация базы данных
def init_db():
    """Инициализирует базу данных для хранения информации о книгах"""
    conn = sqlite3.connect('books.db')
    cursor = conn.cursor()
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS books (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            author TEXT NOT NULL,
            original_filename TEXT NOT NULL,
            site_path TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            chapters_count INTEGER DEFAULT 0
        )
    ''')
    
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS notes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            book_id INTEGER NOT NULL,
            chapter_title TEXT NOT NULL,
            selected_text TEXT NOT NULL,
            note_text TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (book_id) REFERENCES books (id) ON DELETE CASCADE
        )
    ''')
    
    conn.commit()
    conn.close()

# Маршруты Flask
@app.route('/login', methods=['GET', 'POST'])
def login():
    """Страница входа в систему"""
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        if username == USERNAME and password == PASSWORD:
            user = User(username)
            login_user(user)
            next_page = request.args.get('next')
            return redirect(next_page) if next_page else redirect(url_for('index'))
        else:
            flash('Неверный логин или пароль', 'error')
    
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    """Выход из системы"""
    logout_user()
    flash('Вы успешно вышли из системы', 'info')
    return redirect(url_for('login'))

@app.route('/')
@login_required
def index():
    """Главная страница с каталогом книг"""
    # Получаем список всех книг из базы данных
    conn = sqlite3.connect('books.db')
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM books ORDER BY created_at DESC')
    books = cursor.fetchall()
    conn.close()
    
    return render_template('index.html', books=books)

@app.route('/upload', methods=['GET', 'POST'])
@login_required
def upload_book():
    """Страница загрузки новой книги"""
    if request.method == 'POST':
        # Проверяем, что файл был загружен
        if 'file' not in request.files:
            return jsonify({'error': 'Файл не выбран'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Файл не выбран'}), 400
        
        if file and file.filename.lower().endswith('.epub'):
            # Сохраняем загруженный файл
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            
            # Обрабатываем EPUB файл
            processor = EPUBProcessor()
            if processor.load_epub(file_path):
                # Создаем веб-сайт книги
                site_path = processor.create_website(app.config['BOOKS_FOLDER'])
                
                if site_path:
                    # Сохраняем информацию в базу данных
                    conn = sqlite3.connect('books.db')
                    cursor = conn.cursor()
                    # Получаем относительный путь от папки books
                    relative_site_path = os.path.relpath(site_path, app.config['BOOKS_FOLDER']).replace('\\', '/')
                    cursor.execute('''
                        INSERT INTO books (title, author, original_filename, site_path, chapters_count)
                        VALUES (?, ?, ?, ?, ?)
                    ''', (processor.book_title, processor.book_author, filename, 
                          relative_site_path, len(processor.chapters)))
                    book_id = cursor.lastrowid
                    conn.commit()
                    conn.close()
                    
                    # Удаляем временный файл
                    os.remove(file_path)
                    
                    return jsonify({
                        'success': True,
                        'book_id': book_id,
                        'title': processor.book_title,
                        'author': processor.book_author,
                        'chapters_count': len(processor.chapters),
                        'site_path': relative_site_path
                    })
                else:
                    # Удаляем временный файл в случае ошибки создания сайта
                    os.remove(file_path)
                    return jsonify({'error': 'Ошибка при создании сайта книги'}), 500
            else:
                # Удаляем временный файл в случае ошибки обработки EPUB
                os.remove(file_path)
                return jsonify({'error': 'Ошибка при обработке EPUB файла'}), 500
        else:
            return jsonify({'error': 'Недопустимый формат файла. Поддерживается только EPUB'}), 400
    
    return render_template('upload.html')

@app.route('/book/<path:book_path>')
@login_required
def view_book(book_path):
    """Просмотр конкретной книги"""
    return send_from_directory(app.config['BOOKS_FOLDER'], book_path)

@app.route('/book/<path:book_path>/<path:filename>')
@login_required
def book_file(book_path, filename):
    """Обслуживание файлов книги (HTML, CSS, изображения)"""
    full_path = os.path.join(app.config['BOOKS_FOLDER'], book_path)
    return send_from_directory(full_path, filename)

@app.route('/delete/<int:book_id>', methods=['POST'])
@login_required
def delete_book(book_id):
    """Удаление книги"""
    conn = sqlite3.connect('books.db')
    cursor = conn.cursor()
    cursor.execute('SELECT site_path FROM books WHERE id = ?', (book_id,))
    result = cursor.fetchone()
    
    if result:
        site_path = result[0]
        # Удаляем папку с сайтом книги
        import shutil
        full_site_path = os.path.join(app.config['BOOKS_FOLDER'], site_path)
        if os.path.exists(full_site_path):
            shutil.rmtree(full_site_path)
        
        # Удаляем запись из базы данных
        cursor.execute('DELETE FROM books WHERE id = ?', (book_id,))
        conn.commit()
        conn.close()
        
        return jsonify({'success': True})
    else:
        conn.close()
        return jsonify({'error': 'Книга не найдена'}), 404

@app.route('/notes/<int:book_id>')
@login_required
def view_notes(book_id):
    """Просмотр заметок для книги"""
    conn = sqlite3.connect('books.db')
    cursor = conn.cursor()
    
    # Получаем информацию о книге
    cursor.execute('SELECT title, author FROM books WHERE id = ?', (book_id,))
    book = cursor.fetchone()
    
    if not book:
        conn.close()
        return redirect(url_for('index'))
    
    # Получаем все заметки для книги
    cursor.execute('''
        SELECT id, chapter_title, selected_text, note_text, created_at 
        FROM notes 
        WHERE book_id = ? 
        ORDER BY created_at DESC
    ''', (book_id,))
    notes = cursor.fetchall()
    conn.close()
    
    return render_template('notes.html', book=book, notes=notes, book_id=book_id)

@app.route('/api/notes', methods=['POST'])
@login_required
def save_note():
    """Сохранение новой заметки"""
    data = request.get_json()
    
    if not data or not all(k in data for k in ('book_id', 'chapter_title', 'selected_text')):
        return jsonify({'error': 'Недостаточно данных'}), 400
    
    conn = sqlite3.connect('books.db')
    cursor = conn.cursor()
    
    cursor.execute('''
        INSERT INTO notes (book_id, chapter_title, selected_text, note_text)
        VALUES (?, ?, ?, ?)
    ''', (data['book_id'], data['chapter_title'], data['selected_text'], data.get('note_text', '')))
    
    note_id = cursor.lastrowid
    conn.commit()
    conn.close()
    
    return jsonify({'success': True, 'note_id': note_id})

@app.route('/api/notes/<int:note_id>', methods=['PUT'])
@login_required
def update_note(note_id):
    """Обновление заметки"""
    data = request.get_json()
    
    if not data or 'note_text' not in data:
        return jsonify({'error': 'Недостаточно данных'}), 400
    
    conn = sqlite3.connect('books.db')
    cursor = conn.cursor()
    
    cursor.execute('UPDATE notes SET note_text = ? WHERE id = ?', (data['note_text'], note_id))
    
    if cursor.rowcount == 0:
        conn.close()
        return jsonify({'error': 'Заметка не найдена'}), 404
    
    conn.commit()
    conn.close()
    
    return jsonify({'success': True})

@app.route('/api/notes/<int:note_id>', methods=['DELETE'])
@login_required
def delete_note(note_id):
    """Удаление заметки"""
    conn = sqlite3.connect('books.db')
    cursor = conn.cursor()
    
    cursor.execute('DELETE FROM notes WHERE id = ?', (note_id,))
    
    if cursor.rowcount == 0:
        conn.close()
        return jsonify({'error': 'Заметка не найдена'}), 404
    
    conn.commit()
    conn.close()
    
    return jsonify({'success': True})

@app.route('/api/book-info')
@login_required
def get_book_info():
    """Получение информации о книге по пути"""
    book_path = request.args.get('path')
    if not book_path:
        return jsonify({'error': 'Не указан путь к книге'}), 400
    
    conn = sqlite3.connect('books.db')
    cursor = conn.cursor()
    cursor.execute('SELECT id, title, author FROM books WHERE site_path = ?', (book_path,))
    result = cursor.fetchone()
    conn.close()
    
    if result:
        return jsonify({
            'book_id': result[0],
            'title': result[1],
            'author': result[2]
        })
    else:
        return jsonify({'error': 'Книга не найдена'}), 404

@app.route('/download-chapter/<int:book_id>/<int:chapter_index>/<format>')
@login_required
def download_chapter(book_id, chapter_index, format):
    """Скачивание главы в указанном формате"""
    import re
    
    if format not in ['epub', 'docx']:
        return jsonify({'error': 'Неподдерживаемый формат'}), 400
    
    # Получаем информацию о книге
    conn = sqlite3.connect('books.db')
    cursor = conn.cursor()
    cursor.execute('SELECT title, author, site_path FROM books WHERE id = ?', (book_id,))
    result = cursor.fetchone()
    conn.close()
    
    if not result:
        return jsonify({'error': 'Книга не найдена'}), 404
    
    book_title, book_author, site_path = result
    
    # Находим оригинальный EPUB файл
    epub_files = []
    for file in os.listdir(app.config['BOOKS_FOLDER']):
        if file.endswith('.epub'):
            epub_files.append(file)
    
    if not epub_files:
        return jsonify({'error': 'Оригинальный EPUB файл не найден'}), 404
    
    # Берем первый найденный EPUB (можно улучшить логику)
    epub_path = os.path.join(app.config['BOOKS_FOLDER'], epub_files[0])
    
    # Загружаем EPUB и экспортируем главу
    processor = EPUBProcessor()
    if not processor.load_epub(epub_path):
        return jsonify({'error': 'Ошибка при загрузке EPUB файла'}), 500
    
    if chapter_index >= len(processor.chapters):
        return jsonify({'error': 'Глава не найдена'}), 404
    
    chapter_title = processor.chapters[chapter_index]['title']
    safe_title = re.sub(r'[<>:"/\\|?*]', '', chapter_title)
    
    if format == 'docx':
        buffer = processor.export_chapter_to_docx(chapter_index, book_title, book_author)
        if buffer is None:
            return jsonify({'error': 'Ошибка при создании DOCX файла'}), 500
        
        filename = f"{safe_title}.docx"
        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    elif format == 'epub':
        buffer = processor.export_chapter_to_epub(chapter_index, book_title, book_author)
        if buffer is None:
            return jsonify({'error': 'Ошибка при создании EPUB файла'}), 500
        
        filename = f"{safe_title}.epub"
        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/epub+zip'
        )

if __name__ == '__main__':
    # Инициализируем базу данных
    init_db()
    
    # Запускаем приложение
    print("=== Запуск EPUB Cutter Web App ===")
    print("Откройте браузер и перейдите по адресу: http://localhost:5000")
    print("Загружайте EPUB файлы и создавайте веб-сайты книг!")
    
    app.run(debug=True, host='0.0.0.0', port=5000)