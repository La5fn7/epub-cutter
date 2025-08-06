#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Модуль для обработки EPUB файлов
"""

import os
import zipfile
import logging
import tempfile
import io
from pathlib import Path
from xml.etree import ElementTree as ET
import re
import html

logger = logging.getLogger(__name__)


def extract_epub_metadata(opf_root):
    """Извлекает метаданные книги из OPF файла"""
    # Пространства имен для OPF
    namespaces = {
        'opf': 'http://www.idpf.org/2007/opf',
        'dc': 'http://purl.org/dc/elements/1.1/'
    }
    
    # Извлечение названия книги
    title_elem = opf_root.find('.//dc:title', namespaces)
    title = title_elem.text if title_elem is not None and title_elem.text else "Без названия"
    
    # Извлечение автора
    author_elem = opf_root.find('.//dc:creator', namespaces)
    author = author_elem.text if author_elem is not None and author_elem.text else "Неизвестный автор"
    
    return title.strip(), author.strip()


class EPUBProcessor:
    """Класс для обработки EPUB файлов"""
    
    def __init__(self):
        self.book_title = ""
        self.book_author = ""
        self.chapters = []
        self.images = {}
        self.file_mapping = {}
    
    def load_epub(self, epub_path):
        """Загружает EPUB файл и извлекает главы"""
        try:
            self.chapters.clear()
            self.images.clear()
            
            with zipfile.ZipFile(epub_path, 'r') as epub_zip:
                # Чтение структуры EPUB
                container_xml = epub_zip.read('META-INF/container.xml')
                container_root = ET.fromstring(container_xml)
                
                opf_path = container_root.find('.//{urn:oasis:names:tc:opendocument:xmlns:container}rootfile').get('full-path')
                opf_content = epub_zip.read(opf_path)
                opf_root = ET.fromstring(opf_content)
                
                # Извлечение метаданных книги
                self.book_title, self.book_author = extract_epub_metadata(opf_root)
                
                # Загрузка изображений
                opf_dir = os.path.dirname(opf_path)
                for item in opf_root.findall('.//{http://www.idpf.org/2007/opf}item'):
                    media_type = item.get('media-type', '')
                    if media_type.startswith('image/'):
                        href = item.get('href')
                        if href:
                            img_path = os.path.join(opf_dir, href).replace('\\', '/')
                            try:
                                img_data = epub_zip.read(img_path)
                                self.images[img_path] = img_data
                            except Exception as e:
                                logger.debug(f"Ошибка при загрузке изображения {img_path}: {e}")
                
                # Загрузка глав
                spine_items = opf_root.findall('.//{http://www.idpf.org/2007/opf}itemref')
                manifest_items = {item.get('id'): item for item in opf_root.findall('.//{http://www.idpf.org/2007/opf}item')}
                
                # Создаем маппинг файлов для будущих ссылок
                self.file_mapping = {}
                
                for i, spine_item in enumerate(spine_items):
                    idref = spine_item.get('idref')
                    if idref in manifest_items:
                        manifest_item = manifest_items[idref]
                        href = manifest_item.get('href')
                        media_type = manifest_item.get('media-type')
                        
                        if media_type == 'application/xhtml+xml':
                            file_path = os.path.join(opf_dir, href).replace('\\', '/')
                            
                            try:
                                chapter_content = epub_zip.read(file_path)
                                content_str = chapter_content.decode('utf-8')
                                
                                # Сначала пытаемся извлечь заголовок из raw текста
                                title_match = re.search(r'<title[^>]*>([^<]+)</title>', content_str, re.IGNORECASE)
                                title = title_match.group(1).strip() if title_match else f"Глава {i+1}"
                                
                                # Пытаемся исправить некорректные HTML комментарии для корректного XML
                                try:
                                    # Исправляем проблемные комментарии
                                    fixed_content = re.sub(r'<!--\[endif\]---->', '<!--[endif]-->', content_str)
                                    fixed_content = re.sub(r'<!--([^>]*?)---->', r'<!--\1-->', fixed_content)
                                    chapter_root = ET.fromstring(fixed_content.encode('utf-8'))
                                    
                                    # Если XML парсится успешно, используем исправленный контент
                                    chapter_content = fixed_content.encode('utf-8')
                                except Exception as e:
                                    # Если исправление не помогло, используем оригинальный контент
                                    logger.debug(f"XML парсинг не удался для {file_path}: {e}")
                                
                                # Дополнительная попытка найти заголовок в заголовках h1-h6, если title пустой
                                if title == f"Глава {i+1}":
                                    try:
                                        if 'chapter_root' in locals():
                                            for tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                                                heading = chapter_root.find(f'.//{{{http://www.w3.org/1999/xhtml}}}{tag}')
                                                if heading is not None and heading.text:
                                                    title = heading.text.strip()
                                                    break
                                        else:
                                            # Ищем заголовки в raw тексте
                                            for tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                                                h_match = re.search(f'<{tag}[^>]*>([^<]+)</{tag}>', content_str, re.IGNORECASE)
                                                if h_match:
                                                    title = h_match.group(1).strip()
                                                    break
                                    except Exception as e:
                                        logger.debug(f"Ошибка при поиске заголовков в {file_path}: {e}")
                                
                                # Очищаем заголовок от недопустимых символов
                                title = re.sub(r'[<>:"/\\|?*]', '', title)
                                
                                self.chapters.append({
                                    'title': title,
                                    'file_path': file_path,
                                    'content': chapter_content
                                })
                                
                                # Сохраняем маппинг для обработки ссылок
                                self.file_mapping[file_path] = len(self.chapters) - 1
                                
                            except Exception as e:
                                logger.error(f"Ошибка при чтении главы {file_path}: {e}")
                                # Даже если возникла критическая ошибка, сохраняем главу
                                try:
                                    raw_content = epub_zip.read(file_path)
                                    title = f"Глава {i+1}"
                                    
                                    self.chapters.append({
                                        'title': title,
                                        'file_path': file_path,
                                        'content': raw_content
                                    })
                                    
                                    # Сохраняем маппинг для обработки ссылок
                                    self.file_mapping[file_path] = len(self.chapters) - 1
                                    logger.info(f"Добавлена глава с дефолтным названием: {title}")
                                except Exception as e:
                                    logger.error(f"Не удалось загрузить главу {file_path}: {e}")
                                    continue
            
            return True
            
        except Exception as e:
            logger.error(f"Ошибка при загрузке EPUB: {e}")
            return False
    
    def create_website(self, output_path):
        """Создает веб-сайт из загруженной книги"""
        try:
            # Создание главной папки сайта
            site_name = re.sub(r'[<>:"/\\|?*]', '', self.book_title)
            site_path = Path(output_path) / site_name
            site_path.mkdir(parents=True, exist_ok=True)
            
            # Создание папки для изображений
            images_path = site_path / "images"
            images_path.mkdir(exist_ok=True)
            
            # CSS стили (из основного приложения)
            css_content = self._get_website_css()
            with open(site_path / "styles.css", 'w', encoding='utf-8') as f:
                f.write(css_content)
            
            # Создание index.html
            self._create_index_html(site_path)
            
            # Создание страниц глав с навигацией
            selected_chapters = [(i, chapter) for i, chapter in enumerate(self.chapters)]
            
            for idx, (i, chapter) in enumerate(selected_chapters):
                chapter_num = i + 1
                safe_title = re.sub(r'[<>:"/\\|?*]', '', chapter['title'])
                filename = f"chapter_{chapter_num:02d}_{safe_title}.html"
                filepath = site_path / filename
                
                # Обработка содержимого главы
                content = chapter['content'].decode('utf-8')
                content = content.replace('xmlns="http://www.w3.org/1999/xhtml"', '')
                
                # Обработка изображений
                content = self._process_images_for_website(content, images_path)
                
                # Обработка внутренних ссылок
                content = self._process_internal_links(content, selected_chapters, idx)
                
                # Создание навигации
                prev_link = ""
                next_link = ""
                
                if idx > 0:
                    prev_chapter = selected_chapters[idx-1][1]
                    prev_safe_title = re.sub(r'[<>:"/\\|?*]', '', prev_chapter['title'])
                    prev_num = selected_chapters[idx-1][0] + 1
                    prev_filename = f"chapter_{prev_num:02d}_{prev_safe_title}.html"
                    prev_link = f'<a href="{prev_filename}" class="nav-button">← Предыдущая</a>'
                
                if idx < len(selected_chapters) - 1:
                    next_chapter = selected_chapters[idx+1][1]
                    next_safe_title = re.sub(r'[<>:"/\\|?*]', '', next_chapter['title'])
                    next_num = selected_chapters[idx+1][0] + 1
                    next_filename = f"chapter_{next_num:02d}_{next_safe_title}.html"
                    next_link = f'<a href="{next_filename}" class="nav-button">Следующая →</a>'
                
                # Создание полного HTML документа главы
                chapter_html = self._create_chapter_html(chapter, prev_link, next_link, content)
                
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(chapter_html)
            
            return str(site_path)
            
        except Exception as e:
            logger.error(f"Ошибка при создании сайта: {e}")
            return None
    
    def _create_chapter_html(self, chapter, prev_link, next_link, content):
        """Создает HTML для отдельной главы"""
        return f"""<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{html.escape(chapter['title'])} - {html.escape(self.book_title)}</title>
    <link rel="stylesheet" href="styles.css">
</head>
<body>
    <!-- Фиксированная панель инструментов -->
    <div class="toolbar">
        <div class="toolbar-left">
            <h1 class="book-title">{html.escape(self.book_title)}</h1>
        </div>
        <div class="toolbar-right">
            <div class="font-controls">
                <button class="font-btn font-decrease" onclick="changeFontSize(-1)" title="Уменьшить шрифт">A-</button>
                <button class="font-btn font-increase" onclick="changeFontSize(1)" title="Увеличить шрифт">A+</button>
            </div>
            <div class="theme-controls">
                <button class="theme-btn light-theme active" onclick="setTheme('light')" title="Светлая тема">☀️</button>
                <button class="theme-btn dark-theme" onclick="setTheme('dark')" title="Тёмная тема">🌙</button>
            </div>
        </div>
    </div>
    
    <div class="main-content">
        <div class="navigation">
            <a href="index.html" class="nav-button">📚 Содержание</a>
            {prev_link}
            {next_link}
        </div>
        
        <div class="chapter-content" data-chapter-title="{html.escape(chapter['title'])}">
            <h2 class="chapter-title">{html.escape(chapter['title'])}</h2>
            {content}
        </div>
        
        <div class="navigation">
            <a href="index.html" class="nav-button">📚 Содержание</a>
            {prev_link}
            {next_link}
        </div>
    </div>
    
    <script src="/static/js/notes.js"></script>
    <script>
        // Функция изменения размера шрифта
        function changeFontSize(change) {{
            const root = document.documentElement;
            const currentSize = parseInt(getComputedStyle(root).getPropertyValue('--font-size')) || 18;
            const newSize = Math.max(12, Math.min(32, currentSize + change * 2));
            root.style.setProperty('--font-size', newSize + 'px');
            localStorage.setItem('font-size', newSize);
        }}
        
        // Функция переключения тем
        function setTheme(theme) {{
            document.body.className = document.body.className.replace(/theme-\\w+/g, '');
            document.documentElement.className = document.documentElement.className.replace(/theme-\\w+/g, '');
            
            if (theme === 'dark') {{
                document.body.classList.add('theme-dark');
                document.documentElement.classList.add('theme-dark');
            }}
            
            localStorage.setItem('reading-theme', theme);
            
            document.querySelectorAll('.theme-btn').forEach(btn => btn.classList.remove('active'));
            document.querySelector('.' + theme + '-theme').classList.add('active');
        }}
        
        function loadSavedSettings() {{
            // Загружаем сохранённую тему
            const savedTheme = localStorage.getItem('reading-theme') || 'light';
            setTheme(savedTheme);
            
            // Загружаем сохранённый размер шрифта
            const savedFontSize = localStorage.getItem('font-size') || 18;
            document.documentElement.style.setProperty('--font-size', savedFontSize + 'px');
        }}
        
        document.addEventListener('DOMContentLoaded', function() {{
            loadSavedSettings();
            
            const pathParts = window.location.pathname.split('/');
            if (pathParts[1] === 'book' && pathParts[2]) {{
                const bookPath = pathParts[2];
                fetch(`/api/book-info?path=${{encodeURIComponent(bookPath)}}`)
                    .then(response => response.json())
                    .then(data => {{
                        if (data.book_id && window.notesSystem) {{
                            window.notesSystem.bookId = data.book_id;
                            window.notesSystem.addNotesButtonToNavigation();
                        }}
                    }})
                    .catch(error => console.log('Could not load book info:', error));
            }}
        }});
    </script>
</body>
</html>"""
    
    def _create_index_html(self, site_path):
        """Создает главную страницу с оглавлением"""
        selected_chapters = [(i, chapter) for i, chapter in enumerate(self.chapters)]
        
        # Создаем словарь соответствий для ссылок в оглавлении
        self._create_chapter_mapping(selected_chapters)
        
        html_content = f"""<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{self.book_title}</title>
    <link rel="stylesheet" href="styles.css">
</head>
<body>
    <!-- Фиксированная панель инструментов -->
    <div class="toolbar">
        <div class="toolbar-left">
            <h1 class="book-title">{html.escape(self.book_title)}</h1>
        </div>
        <div class="toolbar-right">
            <div class="font-controls">
                <button class="font-btn font-decrease" onclick="changeFontSize(-1)" title="Уменьшить шрифт">A-</button>
                <button class="font-btn font-increase" onclick="changeFontSize(1)" title="Увеличить шрифт">A+</button>
            </div>
            <div class="theme-controls">
                <button class="theme-btn light-theme active" onclick="setTheme('light')" title="Светлая тема">☀️</button>
                <button class="theme-btn dark-theme" onclick="setTheme('dark')" title="Тёмная тема">🌙</button>
            </div>
        </div>
    </div>
    
    <div class="main-content">
        <div class="book-header">
            <p class="book-author">{html.escape(self.book_author)}</p>
        </div>
        
        <div class="toc">
            <h2>📚 Содержание</h2>
            <ul>
"""
        
        for i, (chapter_idx, chapter) in enumerate(selected_chapters):
            chapter_num = chapter_idx + 1
            safe_title = re.sub(r'[<>:"/\\|?*]', '', chapter['title'])
            filename = f"chapter_{chapter_num:02d}_{safe_title}.html"
            
            html_content += f"""            <li>
                <a href="{filename}">
                    <span>{html.escape(chapter['title'])}</span>
                    <span class="chapter-number">{chapter_num}</span>
                </a>
            </li>
"""
        
        html_content += f"""        </ul>
        </div>
        
        <div class="navigation">
            <p>Эта книга содержит {len(selected_chapters)} глав. Выберите главу из оглавления выше для чтения.</p>
        </div>
    </div>
    
    <script>
        // Функция изменения размера шрифта
        function changeFontSize(change) {{
            const root = document.documentElement;
            const currentSize = parseInt(getComputedStyle(root).getPropertyValue('--font-size')) || 18;
            const newSize = Math.max(12, Math.min(32, currentSize + change * 2));
            root.style.setProperty('--font-size', newSize + 'px');
            localStorage.setItem('font-size', newSize);
        }}
        
        // Функция переключения тем
        function setTheme(theme) {{
            document.body.className = document.body.className.replace(/theme-\\w+/g, '');
            document.documentElement.className = document.documentElement.className.replace(/theme-\\w+/g, '');
            
            if (theme === 'dark') {{
                document.body.classList.add('theme-dark');
                document.documentElement.classList.add('theme-dark');
            }}
            
            localStorage.setItem('reading-theme', theme);
            
            document.querySelectorAll('.theme-btn').forEach(btn => btn.classList.remove('active'));
            document.querySelector('.' + theme + '-theme').classList.add('active');
        }}
        
        function loadSavedSettings() {{
            // Загружаем сохранённую тему
            const savedTheme = localStorage.getItem('reading-theme') || 'light';
            setTheme(savedTheme);
            
            // Загружаем сохранённый размер шрифта
            const savedFontSize = localStorage.getItem('font-size') || 18;
            document.documentElement.style.setProperty('--font-size', savedFontSize + 'px');
        }}
        
        document.addEventListener('DOMContentLoaded', function() {{
            loadSavedSettings();
        }});
    </script>
</body>
</html>"""
        
        with open(site_path / "index.html", 'w', encoding='utf-8') as f:
            f.write(html_content)
    
    def _process_images_for_website(self, content, images_path):
        """Обрабатывает изображения для веб-сайта"""
        img_pattern = r'<img[^>]*src=["\']([^"\']+)["\'][^>]*>'
        
        def replace_img(match):
            img_tag = match.group(0)
            src = match.group(1)
            src_clean = src.replace('../', '').replace('./', '')
            
            # Ищем изображение по имени файла
            for img_path, img_data in self.images.items():
                img_filename = os.path.basename(img_path)
                src_filename = os.path.basename(src_clean)
                
                if img_filename == src_filename:
                    # Сохраняем изображение в папку images
                    img_filepath = images_path / img_filename
                    with open(img_filepath, 'wb') as f:
                        f.write(img_data)
                    
                    # Заменяем путь в HTML
                    relative_path = f"images/{img_filename}"
                    new_img_tag = img_tag.replace(f'src="{src}"', f'src="{relative_path}"')
                    new_img_tag = new_img_tag.replace(f"src='{src}'", f'src="{relative_path}"')
                    return new_img_tag
                    
            return img_tag
        
        return re.sub(img_pattern, replace_img, content)
    
    def _process_internal_links(self, content, selected_chapters, current_idx):
        """Обрабатывает внутренние ссылки между главами"""
        # Создаем словарь соответствий: исходный файл -> новый файл
        link_mapping = {}
        
        for idx, (i, chapter) in enumerate(selected_chapters):
            original_path = chapter['file_path']
            # Получаем имя исходного файла без пути
            original_filename = original_path.split('/')[-1]
            
            # Создаем имя нового файла
            chapter_num = i + 1
            safe_title = re.sub(r'[<>:"/\\|?*]', '', chapter['title'])
            new_filename = f"chapter_{chapter_num:02d}_{safe_title}.html"
            
            # Добавляем различные варианты ссылок
            link_mapping[original_filename] = new_filename
            link_mapping[original_path] = new_filename
            
            # Также обрабатываем относительные пути
            if '../' in original_path:
                relative_path = original_path.replace('../', '')
                link_mapping[relative_path] = new_filename
            
            # Обрабатываем пути без расширения (часто используется в ссылках)
            name_without_ext = original_filename.replace('.xhtml', '').replace('.html', '')
            link_mapping[name_without_ext] = new_filename
            
            # Также добавляем вариант с точкой в начале (относительный путь)
            if not original_filename.startswith('./'):
                link_mapping['./' + original_filename] = new_filename
        
        # Паттерн для поиска ссылок
        link_patterns = [
            r'<a[^>]*href=["\'](.*?)["\'][^>]*>',  # обычные ссылки
            r'href=["\'](.*?)["\']',  # просто href атрибуты
        ]
        
        def replace_link(match):
            full_match = match.group(0)
            href_value = match.group(1)
            
            # Пропускаем внешние ссылки и якоря
            if href_value.startswith(('http:', 'https:', 'mailto:', '#')):
                return full_match
            
            # Разделяем ссылку на файл и якорь
            if '#' in href_value:
                file_part, anchor_part = href_value.split('#', 1)
                anchor = '#' + anchor_part
            else:
                file_part = href_value
                anchor = ''
            
            # Убираем относительные пути
            clean_file = file_part.replace('../', '').replace('./', '')
            
            # Ищем соответствие в нашем словаре
            new_filename = None
            
            # Прямое совпадение
            if clean_file in link_mapping:
                new_filename = link_mapping[clean_file]
            else:
                # Поиск по частичному совпадению имени файла
                for original, new in link_mapping.items():
                    if clean_file in original or original.endswith(clean_file):
                        new_filename = new
                        break
            
            if new_filename:
                new_href = new_filename + anchor
                # Отладочный вывод (можно убрать в продакшене)
                if file_part != new_filename:
                    logger.debug(f"Заменяем ссылку: {href_value} -> {new_href}")
                return full_match.replace(href_value, new_href)
            
            return full_match
        
        # Применяем замены для всех паттернов
        for pattern in link_patterns:
            content = re.sub(pattern, replace_link, content)
        
        return content
    
    def _create_chapter_mapping(self, selected_chapters):
        """Создает маппинг для глав"""
        # Этот метод может быть использован для дополнительной логики
        pass
    
    def _get_website_css(self):
        """Возвращает CSS стили для веб-сайта"""
        return """/* CSS стили для EPUB веб-сайта с улучшенным UI */
:root {
    --font-size: 18px;
    --line-height: 1.6;
    --bg-color: #ffffff;
    --text-color: #333333;
    --header-bg: #f8f9fa;
    --border-color: #dee2e6;
    --primary-color: #007bff;
    --secondary-color: #6c757d;
}

.theme-dark {
    --bg-color: #1a1a1a;
    --text-color: #e0e0e0;
    --header-bg: #2c2c2c;
    --border-color: #404040;
    --primary-color: #4dabf7;
    --secondary-color: #adb5bd;
}

body {
    font-family: Georgia, 'Times New Roman', serif;
    font-size: var(--font-size);
    line-height: var(--line-height);
    margin: 0;
    padding: 0;
    background-color: var(--bg-color);
    color: var(--text-color);
    transition: all 0.3s ease;
}

/* Фиксированная панель инструментов */
.toolbar {
    position: fixed;
    top: 0;
    left: 0;
    right: 0;
    background: var(--header-bg);
    border-bottom: 1px solid var(--border-color);
    padding: 10px 20px;
    z-index: 1000;
    display: flex;
    justify-content: space-between;
    align-items: center;
    box-shadow: 0 2px 4px rgba(0,0,0,0.1);
}

.toolbar-left, .toolbar-right {
    display: flex;
    align-items: center;
    gap: 15px;
}

.book-title-mini {
    font-weight: 600;
    color: var(--text-color);
    font-size: 16px;
    margin: 0;
}

.theme-toggle {
    display: flex;
    gap: 5px;
    align-items: center;
}

.theme-btn {
    width: 30px;
    height: 30px;
    border: 2px solid var(--border-color);
    border-radius: 50%;
    cursor: pointer;
    transition: all 0.2s ease;
    display: flex;
    align-items: center;
    justify-content: center;
}

.theme-btn.light {
    background: #ffffff;
    color: #333333;
}

.theme-btn.dark {
    background: #1a1a1a;
    color: #ffffff;
}

.theme-btn.active {
    border-color: var(--primary-color);
    box-shadow: 0 0 0 2px rgba(0, 123, 255, 0.2);
}

.font-controls {
    display: flex;
    align-items: center;
    gap: 10px;
}

.font-btn {
    background: var(--primary-color);
    color: white;
    border: none;
    border-radius: 4px;
    padding: 5px 10px;
    cursor: pointer;
    font-size: 14px;
    transition: all 0.2s ease;
}

.font-btn:hover {
    opacity: 0.8;
}

.font-size-display {
    font-size: 14px;
    color: var(--text-color);
    min-width: 40px;
    text-align: center;
}

/* Основной контент с отступом от панели инструментов */
.main-content {
    margin-top: 70px;
    padding: 20px;
    max-width: 800px;
    margin-left: auto;
    margin-right: auto;
}

.book-header {
    text-align: center;
    margin-bottom: 2rem;
    padding: 1.5rem;
    border-bottom: 1px solid var(--border-color);
}

.book-title {
    font-size: 2.5rem;
    color: var(--text-color);
    margin-bottom: 0.5rem;
}

.book-author {
    font-size: 1.2rem;
    color: var(--secondary-color);
    font-style: italic;
}

.navigation {
    text-align: center;
    margin: 2rem 0;
    padding: 1rem;
}

.nav-button {
    display: inline-block;
    margin: 0 0.5rem;
    padding: 0.75rem 1.5rem;
    background-color: var(--primary-color);
    color: white;
    text-decoration: none;
    border-radius: 5px;
    transition: all 0.3s ease;
    border: none;
    cursor: pointer;
}

.nav-button:hover {
    opacity: 0.8;
    transform: translateY(-1px);
}

.chapter-content {
    background: var(--bg-color);
    padding: 2rem;
    border-radius: 8px;
    line-height: var(--line-height);
}

.chapter-content p {
    margin-bottom: 1em;
    text-align: justify;
}

.chapter-content h1, 
.chapter-content h2, 
.chapter-content h3, 
.chapter-content h4, 
.chapter-content h5, 
.chapter-content h6 {
    color: var(--text-color);
    margin-top: 1.5em;
    margin-bottom: 0.5em;
    line-height: 1.2;
}

.toc {
    background: var(--bg-color);
    padding: 2rem;
    border-radius: 8px;
    border: 1px solid var(--border-color);
}

.toc h2 {
    color: var(--text-color);
    text-align: center;
    margin-bottom: 1.5rem;
}

.toc ul {
    list-style: none;
    padding: 0;
}

.toc li {
    margin-bottom: 0.5rem;
}

.toc a {
    display: flex;
    justify-content: space-between;
    align-items: center;
    padding: 0.75rem 1rem;
    text-decoration: none;
    color: var(--text-color);
    border: 1px solid var(--border-color);
    border-radius: 5px;
    transition: all 0.3s;
}

.toc a:hover {
    background-color: var(--header-bg);
    border-color: var(--primary-color);
}

.chapter-number {
    background-color: var(--primary-color);
    color: white;
    padding: 0.25rem 0.5rem;
    border-radius: 3px;
    font-size: 0.9rem;
}

/* Адаптивность */
@media (max-width: 768px) {
    .toolbar {
        padding: 8px 10px;
        flex-wrap: wrap;
        gap: 10px;
    }
    
    .toolbar-left, .toolbar-right {
        gap: 10px;
    }
    
    .book-title-mini {
        font-size: 14px;
        max-width: 150px;
        overflow: hidden;
        text-overflow: ellipsis;
        white-space: nowrap;
    }
    
    .font-controls {
        gap: 8px;
    }
    
    .main-content {
        margin-top: 80px;
        padding: 15px;
    }
    
    .chapter-content {
        padding: 1rem;
    }
    
    .nav-button {
        padding: 0.5rem 1rem;
        font-size: 0.9rem;
        margin: 0.25rem;
        display: block;
        width: 100%;
        text-align: center;
    }
    
    .navigation {
        margin: 1rem 0;
    }
}

@media (max-width: 480px) {
    .toolbar {
        flex-direction: column;
        padding: 10px;
    }
    
    .toolbar-left, .toolbar-right {
        width: 100%;
        justify-content: space-between;
    }
    
    .main-content {
        margin-top: 100px;
        padding: 10px;
    }
    
    .book-title {
        font-size: 1.8rem;
    }
    
    .theme-btn {
        width: 28px;
        height: 28px;
    }
    
    .font-btn {
        padding: 4px 8px;
        font-size: 12px;
    }
}

/* Стили для кнопок в панели инструментов */
.theme-btn, .font-btn {
    background: var(--bg-color);
    border: 1px solid var(--border-color);
    color: var(--text-color);
    cursor: pointer;
    transition: all 0.2s ease;
    font-family: inherit;
    outline: none;
}

.theme-btn:hover, .font-btn:hover {
    background: var(--primary-color);
    color: white;
    transform: scale(1.05);
}

.theme-btn.active {
    background: var(--primary-color);
    color: white;
    border-color: var(--primary-color);
}

.font-btn {
    padding: 6px 12px;
    border-radius: 4px;
    font-size: 14px;
    font-weight: bold;
    min-width: 36px;
}

.theme-btn {
    padding: 6px 10px;
    border-radius: 4px;
    font-size: 16px;
    min-width: 40px;
}

/* Темная тема */
.theme-dark {
    background-color: #1a1a1a;
    color: #e0e0e0;
}

.theme-dark .toolbar {
    background: var(--header-bg);
    border-bottom-color: var(--border-color);
}

.theme-dark .book-header {
    border-bottom-color: #444;
}

.theme-dark .book-title {
    color: var(--text-color);
}

.theme-dark .chapter-content,
.theme-dark .toc {
    background: var(--bg-color);
    color: var(--text-color);
}

.theme-dark .nav-button {
    background-color: var(--primary-color);
    color: white;
    border-color: var(--primary-color);
}

.theme-dark .nav-button:hover {
    background-color: #3a8dff;
}

.theme-dark .toc a {
    color: var(--text-color);
    border-color: var(--border-color);
}

.theme-dark .toc a:hover {
    background-color: #3c3c3c;
    border-color: var(--primary-color);
}

.theme-dark .chapter-title {
    color: var(--text-color);
    border-bottom-color: var(--border-color);
}"""
    
    def export_chapter_to_docx(self, chapter_index, book_title, book_author):
        """Экспортирует главу в формат DOCX"""
        try:
            from docx import Document
            from docx.shared import Inches
            from html import unescape
            
            if chapter_index >= len(self.chapters):
                return None
                
            chapter = self.chapters[chapter_index]
            
            # Создаем новый документ
            doc = Document()
            
            # Добавляем заголовок книги
            title = doc.add_heading(book_title, 0)
            author = doc.add_paragraph(f'Автор: {book_author}')
            author.alignment = 1  # Центрирование
            
            # Добавляем заголовок главы
            chapter_title = doc.add_heading(chapter['title'], level=1)
            
            # Обрабатываем содержимое главы
            content = chapter['content']
            if isinstance(content, bytes):
                content = content.decode('utf-8')
            
            # Парсим HTML для извлечения текста
            try:
                # Исправляем комментарии если нужно
                content = re.sub(r'<!--\[endif\]---->', '<!--[endif]-->', content)
                content = re.sub(r'<!--([^>]*?)---->', r'<!--\1-->', content)
                
                root = ET.fromstring(content)
                body = root.find('.//{http://www.w3.org/1999/xhtml}body')
                
                if body is not None:
                    self._process_html_to_docx(body, doc)
                else:
                    # Если не удалось парсить, извлекаем текст регулярными выражениями
                    self._extract_text_to_docx(content, doc)
                    
            except Exception as e:
                # Fallback - извлекаем текст простыми методами
                logger.debug(f"Ошибка при парсинге HTML для DOCX: {e}")
                self._extract_text_to_docx(content, doc)
            
            # Сохраняем в память
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            
            return docx_buffer
            
        except Exception as e:
            logger.error(f"Ошибка при экспорте в DOCX: {e}")
            return None
    
    def _process_html_to_docx(self, element, doc):
        """Обрабатывает HTML элементы и добавляет в DOCX документ"""
        for child in element:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            
            if tag in ['p']:
                if child.text or len(child) > 0:
                    paragraph = doc.add_paragraph()
                    self._add_text_with_formatting(child, paragraph)
            elif tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(tag[1]) + 1  # h1 -> level 2, h2 -> level 3, etc.
                if child.text:
                    doc.add_heading(child.text.strip(), level=level)
            elif tag == 'img':
                self._add_image_to_docx(child, doc)
            elif tag == 'div':
                self._process_html_to_docx(child, doc)
    
    def _add_text_with_formatting(self, element, paragraph):
        """Добавляет текст с форматированием в параграф"""
        if element.text:
            paragraph.add_run(element.text)
        
        for child in element:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            
            if tag == 'em' or tag == 'i':
                run = paragraph.add_run(child.text or '')
                run.italic = True
            elif tag == 'strong' or tag == 'b':
                run = paragraph.add_run(child.text or '')
                run.bold = True
            elif tag == 'img':
                # Для изображений внутри параграфов добавляем только ссылку
                src = child.get('src', '')
                img_name = src.split('/')[-1] if src else 'изображение'
                paragraph.add_run(f" [Изображение: {img_name}] ")
            else:
                paragraph.add_run(child.text or '')
            
            if child.tail:
                paragraph.add_run(child.tail)
    
    def _add_image_to_docx(self, img_element, doc):
        """Добавляет изображение в DOCX документ"""
        try:
            from docx.shared import Inches
            
            # Получаем src атрибут
            src = img_element.get('src')
            if not src:
                return
            
            # Ищем изображение в загруженных изображениях
            img_data = None
            for img_path, data in self.images.items():
                if src in img_path or os.path.basename(src) == os.path.basename(img_path):
                    img_data = data
                    break
            
            if img_data:
                # Создаем новый параграф для изображения
                paragraph = doc.add_paragraph()
                paragraph.alignment = 1  # Центрирование
                
                # Добавляем изображение
                img_stream = io.BytesIO(img_data)
                try:
                    run = paragraph.add_run()
                    run.add_picture(img_stream, width=Inches(4))
                except Exception as e:
                    # Если не получилось добавить изображение, добавляем текст
                    logger.debug(f"Не удалось добавить изображение {src} в DOCX: {e}")
                    run = paragraph.add_run()
                    run.text = f"[Изображение: {src.split('/')[-1]}]"
                    
        except Exception as e:
            logger.warning(f"Не удалось добавить изображение в DOCX: {e}")

    def _extract_text_to_docx(self, html_content, doc):
        """Извлекает текст из HTML простыми методами"""
        from html import unescape
        
        # Удаляем скрипты и стили
        html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Извлекаем заголовки
        headings = re.findall(r'<h[1-6][^>]*>(.*?)</h[1-6]>', html_content, re.DOTALL | re.IGNORECASE)
        for heading in headings:
            clean_heading = re.sub(r'<[^>]+>', '', heading)
            clean_heading = unescape(clean_heading).strip()
            if clean_heading:
                doc.add_heading(clean_heading, level=2)
        
        # Извлекаем параграфы
        paragraphs = re.findall(r'<p[^>]*>(.*?)</p>', html_content, re.DOTALL | re.IGNORECASE)
        for para in paragraphs:
            clean_para = re.sub(r'<[^>]+>', '', para)
            clean_para = unescape(clean_para).strip()
            if clean_para and len(clean_para) > 10:  # Игнорируем очень короткие параграфы
                doc.add_paragraph(clean_para)
    
    def export_chapter_to_epub(self, chapter_index, book_title, book_author):
        """Экспортирует главу в формат EPUB"""
        try:
            if chapter_index >= len(self.chapters):
                return None
                
            chapter = self.chapters[chapter_index]
            
            # Создаем временный EPUB файл
            epub_buffer = io.BytesIO()
            
            with zipfile.ZipFile(epub_buffer, 'w', zipfile.ZIP_DEFLATED) as epub_zip:
                # mimetype (должен быть первым и несжатым)
                epub_zip.writestr('mimetype', 'application/epub+zip', zipfile.ZIP_STORED)
                
                # META-INF/container.xml
                container_xml = '''<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
    <rootfiles>
        <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
    </rootfiles>
</container>'''
                epub_zip.writestr('META-INF/container.xml', container_xml)
                
                # OEBPS/content.opf
                import uuid
                book_id = str(uuid.uuid4())
                
                content_opf = f'''<?xml version="1.0" encoding="utf-8"?>
<package xmlns="http://www.idpf.org/2007/opf" version="2.0" unique-identifier="BookId">
    <metadata xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:opf="http://www.idpf.org/2007/opf">
        <dc:identifier id="BookId" opf:scheme="UUID">{book_id}</dc:identifier>
        <dc:title>{html.escape(book_title)}</dc:title>
        <dc:creator opf:role="aut">{html.escape(book_author)}</dc:creator>
        <dc:language>ru</dc:language>
        <dc:date>{html.escape(str(__import__('datetime').datetime.now().strftime('%Y-%m-%d')))}</dc:date>
    </metadata>
    <manifest>
        <item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>
        <item id="chapter" href="chapter.xhtml" media-type="application/xhtml+xml"/>
    </manifest>
    <spine toc="ncx">
        <itemref idref="chapter"/>
    </spine>
</package>'''
                epub_zip.writestr('OEBPS/content.opf', content_opf)
                
                # OEBPS/toc.ncx
                toc_ncx = f'''<?xml version="1.0" encoding="utf-8"?>
<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">
    <head>
        <meta name="dtb:uid" content="{book_id}"/>
        <meta name="dtb:depth" content="1"/>
        <meta name="dtb:totalPageCount" content="0"/>
        <meta name="dtb:maxPageNumber" content="0"/>
    </head>
    <docTitle>
        <text>{html.escape(book_title)}</text>
    </docTitle>
    <navMap>
        <navPoint id="navPoint-1" playOrder="1">
            <navLabel>
                <text>{html.escape(chapter['title'])}</text>
            </navLabel>
            <content src="chapter.xhtml"/>
        </navPoint>
    </navMap>
</ncx>'''
                epub_zip.writestr('OEBPS/toc.ncx', toc_ncx)
                
                # OEBPS/chapter.xhtml
                content = chapter['content']
                if isinstance(content, bytes):
                    content = content.decode('utf-8')
                
                # Убираем проблемные пространства имен и исправляем структуру
                content = content.replace('xmlns="http://www.w3.org/1999/xhtml"', '')
                content = re.sub(r'<!DOCTYPE[^>]*>', '', content)
                
                chapter_xhtml = f'''<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.1//EN" "http://www.w3.org/TR/xhtml11/DTD/xhtml11.dtd">
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
    <title>{html.escape(chapter['title'])}</title>
</head>
<body>
    <h1>{html.escape(chapter['title'])}</h1>
    {content}
</body>
</html>'''
                epub_zip.writestr('OEBPS/chapter.xhtml', chapter_xhtml)
            
            epub_buffer.seek(0)
            return epub_buffer
            
        except Exception as e:
            logger.error(f"Ошибка при экспорте в EPUB: {e}")
            return None